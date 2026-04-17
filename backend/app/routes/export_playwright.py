"""
Playwright-based PDF/DOCX export.

Architecture:
1. Frontend has a dedicated /resume/export/:evaluationId page that renders
   the EXACT SAME template component used in the main UI (but stripped of
   interactive elements like AI buttons and animations).
2. When user clicks "Export PDF", the backend:
   a. Generates a short-lived JWT export token
   b. Launches a headless Chromium via Playwright
   c. Navigates to the frontend export page with the token
   d. Waits for full render + font loading
   e. Calls page.pdf() to generate pixel-perfect PDF
   f. Returns the PDF bytes to the user
3. For DOCX, we convert the PDF using pdf2docx (best fidelity) or
   fall back to a basic python-docx structure.
"""

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse, JSONResponse
from ..middleware import get_current_user
from ..models import UserModel
from ..database import get_db, to_object_id
from ..config import settings
from jose import jwt
import sys
import threading
import concurrent.futures
import traceback
import io
import asyncio
import datetime
import re

router = APIRouter()

# ─── Export Token (short-lived JWT for headless browser auth) ──────────────────

def _create_export_token(user_id: str, evaluation_id: str) -> str:
    """Create a 60-second JWT that the export page uses to fetch data."""
    payload = {
        "sub": user_id,
        "eval_id": evaluation_id,
        "purpose": "export",
        "exp": datetime.datetime.utcnow() + datetime.timedelta(seconds=120),
    }
    return jwt.encode(payload, settings.JWT_SECRET, algorithm=settings.JWT_ALGORITHM)


def _verify_export_token(token: str) -> dict:
    """Verify and decode an export token. Raises on invalid/expired."""
    try:
        payload = jwt.decode(token, settings.JWT_SECRET, algorithms=[settings.JWT_ALGORITHM])
        if payload.get("purpose") != "export":
            raise ValueError("Not an export token")
        return payload
    except Exception as e:
        raise HTTPException(status_code=401, detail=f"Invalid export token: {e}")


# ─── Render Data Endpoint (called by the frontend export page) ────────────────

@router.get("/render-data/{evaluation_id}")
async def get_render_data(evaluation_id: str, export_token: str = Query(...)):
    """
    Returns the final merged resume data for rendering.
    Authenticated via short-lived export token (NOT Bearer JWT).
    This endpoint is called by the frontend ExportPage.tsx.
    """
    payload = _verify_export_token(export_token)
    user_id = payload["sub"]

    db = get_db()

    # 1. Get evaluation (with fallback for legacy analysis_result IDs)
    evaluation = await db.evaluations.find_one({
        "_id": to_object_id(evaluation_id),
        "user_id": user_id,
    })
    
    if not evaluation:
        # Fallback: maybe the frontend passed an old analysis_results ID?
        old_analysis = await db.analysis_results.find_one({
            "_id": to_object_id(evaluation_id),
            "user_id": user_id,
        })
        if old_analysis and old_analysis.get("resume_id"):
            evaluation = await db.evaluations.find_one(
                {"resume_id": old_analysis["resume_id"], "user_id": user_id},
                sort=[("_id", -1)]
            )
            
    if not evaluation:
        raise HTTPException(status_code=404, detail="Evaluation not found")

    resume_id = evaluation.get("resume_id")
    if not resume_id:
        raise HTTPException(status_code=404, detail="No resume linked to this evaluation")

    # 2. Get resume
    resume = await db.resumes.find_one({"_id": to_object_id(resume_id)})
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    # 3. Get base structured data
    base_data = (
        resume.get("optimized_resume_json")
        or resume.get("structured_resume_json")
        or resume.get("extracted_data", {})
    )

    # 4. Merge accepted edits from BOTH resume AND evaluation
    from ..services.resume_merger import merge_accepted_edits, filter_deleted_blocks

    # Start with resume-level edits
    all_edits = dict(resume.get("accepted_edits", {}))
    # Overlay evaluation-specific edits (these are the latest)
    all_edits.update(evaluation.get("accepted_edits", {}))

    merged_data, merge_count = merge_accepted_edits(base_data, all_edits)
    print(f"[Export Render] Merged {merge_count} edits for evaluation {evaluation_id}")

    # 5. Filter deleted blocks from both sources
    all_deleted = list(set(
        resume.get("deleted_blocks", []) +
        evaluation.get("deleted_blocks", [])
    ))
    merged_data, del_count = filter_deleted_blocks(merged_data, all_deleted)
    if del_count > 0:
        print(f"[Export Render] Filtered {del_count} deleted blocks")

    # 6. Get template preference
    template_id = resume.get("selected_template", "modern-ats")

    return {
        "structured_resume": merged_data,
        "template_id": template_id,
        "evaluation_id": evaluation_id,
        "resume_id": resume_id,
        "user_name": evaluation.get("user_name", merged_data.get("name", "resume")),
        "job_title": evaluation.get("job_title", "export"),
    }


# ─── Isolated Proactor Execution (Windows Fix) ────────────────────────────────

async def _run_in_isolated_proactor(func, *args, **kwargs):
    """
    Runs an async function in a dedicated background thread with its own
    ProactorEventLoop. This is non-blocking for the main server thread.
    """
    def _worker():
        if sys.platform == "win32":
            # Force the thread to use Proactor loop for subprocesses
            asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())
        
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            return loop.run_until_complete(func(*args, **kwargs))
        finally:
            loop.close()

    return await asyncio.to_thread(_worker)


# ─── Playwright PDF Generation ────────────────────────────────────────────────

async def _generate_pdf_with_playwright(evaluation_id: str, export_token: str) -> bytes:
    """
    Launch headless Chromium, open the export page, wait for render, 
    and generate a pixel-perfect PDF.
    """
    from playwright.async_api import async_playwright

    frontend_url = settings.FRONTEND_URL or "http://127.0.0.1:5173"
    # Ensure frontend_url uses IPv4 and correct port to avoid Windows resolution mismatches
    frontend_url = frontend_url.replace("localhost", "127.0.0.1")
    export_url = f"{frontend_url}/resume/export/{evaluation_id}?export_token={export_token}"

    print(f"[Playwright] Attempting navigation to target: {export_url}")

    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=True,
            args=[
                '--no-sandbox',
                '--disable-setuid-sandbox',
                '--disable-dev-shm-usage',
                '--font-render-hinting=none',
            ]
        )

        page = await browser.new_page(
            viewport={"width": 794, "height": 1123},  # A4 at 96 DPI
        )

        try:
            # Navigate and wait for base load (increased timeout for reliability)
            await page.goto(export_url, wait_until="load", timeout=60000)

            # Wait for the export root to be marked as ready (this confirms data is fetched)
            await page.wait_for_selector('#export-root[data-ready="true"]', timeout=30000)

            # Ensure all fonts are loaded
            await page.evaluate("() => document.fonts.ready")

            # Small extra delay for any final paints
            await asyncio.sleep(0.5)

            # Generate PDF
            pdf_bytes = await page.pdf(
                format="A4",
                print_background=True,
                prefer_css_page_size=True,
                margin={
                    "top": "0mm",
                    "bottom": "0mm",
                    "left": "0mm",
                    "right": "0mm",
                },
            )

            print(f"[Playwright] PDF generated: {len(pdf_bytes)} bytes")
            return pdf_bytes

        except Exception as e:
            # Capture what went wrong, but only if the page context still exists
            try:
                if not page.is_closed():
                    error_el = await page.query_selector('#export-error')
                    if error_el:
                        error_text = await error_el.text_content()
                        print(f"[Playwright] Export page error: {error_text}")
            except Exception:
                pass # Avoid masking the original error
            raise e
        finally:
            await browser.close()


# ─── PDF Export Endpoint ──────────────────────────────────────────────────────

@router.get("/pdf/{evaluation_id}")
async def export_pdf(evaluation_id: str, current_user: UserModel = Depends(get_current_user)):
    """
    Generates a pixel-perfect PDF using Playwright headless browser.
    The exported PDF is identical to what the user sees on screen.
    """
    try:
        # Generate export token for the headless browser
        export_token = _create_export_token(str(current_user.id), evaluation_id)

        # Generate PDF
        # Use non-blocking isolated proactor thread to avoid deadlock and NotImplementedError
        pdf_bytes = await _run_in_isolated_proactor(
            _generate_pdf_with_playwright, evaluation_id, export_token
        )

        # Build filename
        db = get_db()
        evaluation = await db.evaluations.find_one({
            "_id": to_object_id(evaluation_id),
            "user_id": str(current_user.id),
        })
        if not evaluation:
            # Fallback
            old_analysis = await db.analysis_results.find_one({
                "_id": to_object_id(evaluation_id),
                "user_id": str(current_user.id),
            })
            if old_analysis and old_analysis.get("resume_id"):
                evaluation = await db.evaluations.find_one(
                    {"resume_id": old_analysis["resume_id"]},
                    sort=[("_id", -1)]
                )

        # Fetch resume to get the name from the data if needed
        resume = None
        if evaluation and evaluation.get("resume_id"):
            resume = await db.resumes.find_one({"_id": to_object_id(evaluation["resume_id"])})

        filename = _build_filename(evaluation, current_user, "pdf", resume=resume)
        print(f"[Export PDF] Suggested filename: {filename}")

        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )

    except Exception as e:
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"PDF export failed: {str(e)}"
        )


# ─── DOCX Export Endpoint ─────────────────────────────────────────────────────

@router.get("/docx/{evaluation_id}")
async def export_docx(evaluation_id: str, current_user: UserModel = Depends(get_current_user)):
    """
    Generates a DOCX export. Uses the merged data directly with python-docx
    for the best text-based DOCX fidelity (ATS-readable).
    """
    try:
        db = get_db()
        evaluation = await db.evaluations.find_one({
            "_id": to_object_id(evaluation_id),
            "user_id": str(current_user.id),
        })
        
        if not evaluation:
            # Fallback
            old_analysis = await db.analysis_results.find_one({
                "_id": to_object_id(evaluation_id),
                "user_id": str(current_user.id),
            })
            if old_analysis and old_analysis.get("resume_id"):
                evaluation = await db.evaluations.find_one(
                    {"resume_id": old_analysis["resume_id"]},
                    sort=[("_id", -1)]
                )
                
        if not evaluation:
            raise HTTPException(status_code=404, detail="Evaluation not found (even after fallback)")

        resume_id = evaluation.get("resume_id")
        resume = await db.resumes.find_one({"_id": to_object_id(resume_id)})
        if not resume:
            raise HTTPException(status_code=404, detail="Resume not found")

        # Merge data (same logic as render-data)
        base_data = (
            resume.get("optimized_resume_json")
            or resume.get("structured_resume_json")
            or resume.get("extracted_data", {})
        )

        from ..services.resume_merger import merge_accepted_edits, filter_deleted_blocks

        all_edits = dict(resume.get("accepted_edits", {}))
        all_edits.update(evaluation.get("accepted_edits", {}))

        merged_data, _ = merge_accepted_edits(base_data, all_edits)

        all_deleted = list(set(
            resume.get("deleted_blocks", []) +
            evaluation.get("deleted_blocks", [])
        ))
        merged_data, _ = filter_deleted_blocks(merged_data, all_deleted)

        # Build DOCX
        docx_bytes = _build_docx(merged_data)

        # Fetch resume to get the name from the data if needed
        resume_doc = await db.resumes.find_one({"_id": to_object_id(resume_id)})

        filename = _build_filename(evaluation, current_user, "docx", resume=resume_doc)
        print(f"[Export DOCX] Suggested filename: {filename}")

        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )

    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"DOCX export failed: {str(e)}")


# ─── DOCX Builder (clean, ATS-friendly) ──────────────────────────────────────

def _build_docx(data: dict) -> bytes:
    """Build a clean ATS-friendly DOCX from merged structured data."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Default font style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.2

    def add_heading(title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x37, 0x30, 0xa3)
        run.font.name = 'Georgia'
        # Bottom border
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '10')
        bottom.set(qn('w:color'), '3730A3')
        bottom.set(qn('w:space'), '1')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_bullet(text: str):
        for paragraph in str(text).split('\n'):
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            if paragraph.startswith('-') or paragraph.startswith('•') or paragraph.startswith('\u2022'):
                paragraph = paragraph[1:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.2
            run = p.add_run(paragraph)
            run.font.size = Pt(10.5)
            run.font.name = 'Georgia'

    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(4)
    run = name_p.add_run(data.get("name") or "Your Name")
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Georgia'
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Contact
    contact_parts = []
    if data.get("phone"): contact_parts.append(data["phone"])
    if data.get("email"): contact_parts.append(data["email"])
    links = data.get("links", {})
    if links.get("linkedin"): contact_parts.append("LinkedIn")
    if links.get("github"): contact_parts.append("GitHub")
    if links.get("website"): contact_parts.append("Portfolio")
    if contact_parts:
        c_p = doc.add_paragraph()
        c_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_p.paragraph_format.space_after = Pt(8)
        run = c_p.add_run("  ·  ".join(contact_parts))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.name = 'Georgia'

    # Divider
    div = doc.add_paragraph()
    div.paragraph_format.space_before = Pt(0)
    div.paragraph_format.space_after = Pt(12)
    pPr = div._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '16')
    top.set(qn('w:color'), '3730A3')
    top.set(qn('w:space'), '1')
    pBdr.append(top)
    pPr.append(pBdr)

    # Summary
    if data.get("summary"):
        add_heading("PROFESSIONAL SUMMARY")
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(str(data["summary"]))
        run.font.size = Pt(10.5)
        run.font.name = 'Georgia'

    # Experience
    if data.get("experience"):
        add_heading("EXPERIENCE")
        for item in data["experience"]:
            add_bullet(str(item))

    # Education
    if data.get("education"):
        add_heading("EDUCATION")
        for item in data["education"]:
            add_bullet(str(item))

    # Technical Skills
    skills_cat = data.get("skills_categorized", {})
    if skills_cat and isinstance(skills_cat, dict):
        add_heading("TECHNICAL SKILLS")
        for category, items in skills_cat.items():
            if not items:
                continue
            clean_items = [str(item) for item in items if item]
            if not clean_items:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.2
            cat_run = p.add_run(f"{category}: ")
            cat_run.bold = True
            cat_run.font.size = Pt(10.5)
            cat_run.font.name = 'Georgia'
            val_run = p.add_run(", ".join(clean_items))
            val_run.font.size = Pt(10.5)
            val_run.font.name = 'Georgia'
    elif data.get("skills"):
        add_heading("SKILLS")
        p = doc.add_paragraph()
        p.add_run(" · ".join(str(s) for s in data["skills"]))

    # Remaining sections
    for section_name in ["projects", "certifications", "publications", "volunteering"]:
        items = data.get(section_name, [])
        if items:
            add_heading(section_name.upper())
            for item in items:
                add_bullet(str(item))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── Filename Builder ─────────────────────────────────────────────────────────

def _build_filename(evaluation: dict | None, user: UserModel, ext: str, resume: dict | None = None) -> str:
    """
    Build filename: username_jobtitle.ext
    Lowercase, spaces → underscores, strip special chars.
    """
    # 1. Identify the user name component
    # Priority order:
    #   - eval.user_name
    #   - resume.extracted_data.name
    #   - resume.structured_resume_json.basics.name
    #   - eval.resume_name (minus ext)
    #   - user.name
    #   - user.email (prefix)
    
    name_val = ""
    source = "default"
    
    if evaluation and evaluation.get("user_name"):
        name_val = evaluation["user_name"]
        source = "evaluation.user_name"
    
    if not name_val and resume:
        # Check all possible data buckets in resume
        buckets = [
            resume.get("extracted_data", {}),
            resume.get("structured_resume_json", {}),
            resume.get("optimized_resume_json", {}),
            resume.get("structured_resume_json", {}).get("basics", {}),
        ]
        for b in buckets:
            if isinstance(b, dict) and b.get("name"):
                name_val = b["name"]
                source = "resume_bucket_search"
                break
                
    if not name_val and evaluation and evaluation.get("resume_name"):
        name_val = evaluation["resume_name"].split('.')[0]
        source = "evaluation.resume_name"

    if not name_val:
        name_val = getattr(user, "name", "")
        source = "user.name"
    
    if not name_val and getattr(user, "email", ""):
        name_val = user.email.split('@')[0]
        source = "user.email"

    name_part = name_val or "resume"
    print(f"[Build Filename] FINAL NAME SOURCE: {source} -> '{name_part}'")

    # 2. Identify the job context component
    job_part = "export"
    if evaluation and evaluation.get("job_title"):
        job_part = evaluation["job_title"]

    def clean(s: str) -> str:
        if not s:
            return "unknown"
        s = s.lower().strip()
        # Replace non-filename-safe chars with underscore
        # We allow a-z, 0-9, and underscore. Everything else becomes underscore.
        s = re.sub(r'[^a-z0-9]', '_', s)
        # Collapse multiple underscores
        s = re.sub(r'_+', '_', s)
        return s.strip('_') or "unknown"

    final_name = clean(name_part)
    final_job = clean(job_part)
    
    return f"{final_name}_{final_job}.{ext}"
